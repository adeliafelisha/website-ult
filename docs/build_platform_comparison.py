from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name("perbandingan-wordpress-vs-laravel-filament-ult.docx")
NAVY = "3B145F"
PURPLE = "7E22CE"
LILAC = "F3E8FF"
GOLD = "D99A00"
INK = "202124"
MUTED = "5F6368"
LIGHT = "F5F3F7"
GREEN = "137333"
RED = "B3261E"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 15, 7),
    ("Heading 2", 13, PURPLE, 11, 5),
    ("Heading 3", 11.5, NAVY, 8, 3),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Number"]:
    st = styles[list_name]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.18)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.1

if "Pros" not in styles:
    pros = styles.add_style("Pros", WD_STYLE_TYPE.PARAGRAPH)
    pros.base_style = styles["List Bullet"]
    pros.font.color.rgb = RGBColor.from_string(GREEN)
if "Cons" not in styles:
    cons = styles.add_style("Cons", WD_STYLE_TYPE.PARAGRAPH)
    cons.base_style = styles["List Bullet"]
    cons.font.color.rgb = RGBColor.from_string(RED)

def font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for old in list(grid): grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)

def page_break():
    doc.add_page_break()

def bullet(text, kind="normal"):
    style = "List Bullet" if kind == "normal" else ("Pros" if kind == "pro" else "Cons")
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number"); p.add_run(text); return p

def callout(title, text, fill=LILAC, accent=PURPLE):
    t = doc.add_table(rows=1, cols=1)
    table_geometry(t, [9360])
    c = t.cell(0, 0); shade(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    font(p.add_run(title + "  "), bold=True, color=accent)
    font(p.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def platform_block(title, pros, cons):
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(); font(p.add_run("Kelebihan"), bold=True, color=GREEN)
    for x in pros: bullet(x, "pro")
    p = doc.add_paragraph(); font(p.add_run("Kekurangan"), bold=True, color=RED)
    for x in cons: bullet(x, "con")

# Running furniture
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("ULT UNPAD  |  Kajian Platform Website"), size=8.5, bold=True, color=MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Dokumen kerja untuk evaluasi internal - Agustus 2026"), size=8, color=MUTED)

# Cover / memo masthead
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
font(p.add_run("KAJIAN PERBANDINGAN PLATFORM"), size=10, bold=True, color=GOLD)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
font(p.add_run("WordPress vs Laravel + Filament"), size=25, bold=True, color=NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
font(p.add_run("Website Unit Layanan Terpadu Universitas Padjadjaran"), size=13, color=PURPLE)

meta = [("Tujuan", "Mendukung evaluasi migrasi website lama WordPress ke aplikasi Laravel dengan CMS Filament"),
        ("Ruang lingkup", "Frontend, pengalaman pengguna, CRUD/admin, backend, database, keamanan, operasional, dan pengembangan"),
        ("Basis kajian", "Kondisi aplikasi Laravel ULT per Agustus 2026 dan karakteristik umum WordPress"),
        ("Status", "Dokumen kerja - perlu divalidasi terhadap konfigurasi WordPress/server lama")]
t = doc.add_table(rows=len(meta), cols=2)
table_geometry(t, [1800, 7560])
for i, (k, v) in enumerate(meta):
    shade(t.cell(i,0), LIGHT); font(t.cell(i,0).paragraphs[0].add_run(k), bold=True, color=NAVY)
    font(t.cell(i,1).paragraphs[0].add_run(v), color=INK)

doc.add_heading("Ringkasan eksekutif", level=1)
doc.add_paragraph("WordPress unggul ketika kebutuhan utama adalah menerbitkan konten dengan cepat menggunakan ekosistem tema dan plugin yang matang. Laravel + Filament unggul ketika ULT membutuhkan aplikasi yang lebih terstruktur, alur data khusus, kontrol keamanan yang eksplisit, integrasi sistem, dan kemampuan berkembang melampaui website informasi.")
callout("Kesimpulan utama", "Untuk arah ULT sebagai portal layanan publik yang bilingual, aksesibel, memiliki direktori layanan, artikel, analitik internal, dan CMS terkontrol, Laravel + Filament lebih sesuai sebagai fondasi jangka menengah-panjang. Konsekuensinya adalah kebutuhan tim teknis, deployment, monitoring, backup, dan pemeliharaan kode yang lebih disiplin.")

doc.add_heading("Cara membaca dokumen", level=2)
bullet("WordPress berarti platform website lama berbasis WordPress; detail plugin, tema, versi PHP, dan konfigurasi server lama belum diaudit dalam kajian ini.")
bullet("Laravel + Filament berarti aplikasi baru ULT yang saat ini sudah memiliki frontend bilingual, CMS CRUD, artikel, layanan, FAQ, kontak, tautan cepat, aksesibilitas, dan dashboard trafik internal.")
bullet("Pernyataan keamanan tidak berarti salah satu platform otomatis aman. Keamanan aktual bergantung pada versi, konfigurasi, kode, kredensial, hosting, backup, dan disiplin operasional.")

page_break()
doc.add_heading("1. Matriks ringkas per kategori", level=1)
rows = [
    ("Kecepatan peluncuran", "Cepat dengan tema/plugin", "Lebih lambat di awal; sangat fleksibel setelah fondasi tersedia", "WordPress untuk situs sederhana"),
    ("Tampilan frontend", "Mudah melalui page builder, tetapi rawan tidak konsisten", "Kontrol penuh dan konsisten; membutuhkan developer", "Laravel untuk desain ULT khusus"),
    ("CRUD admin", "Editor konten matang dan familiar", "CRUD domain-spesifik, validasi, role, dan workflow dapat dibentuk", "Seimbang, tergantung kebutuhan"),
    ("Performa", "Baik jika ringan dan tercache; plugin dapat membebani", "Dapat sangat cepat; perlu optimasi query/cache/deployment", "Laravel untuk skala terkontrol"),
    ("Integrasi", "Banyak plugin siap pakai", "API dan integrasi khusus lebih bersih dan teruji", "Laravel untuk integrasi khusus"),
    ("Security surface", "Core matang, tetapi plugin/theme memperluas permukaan serangan", "Permukaan lebih terkontrol, tetapi tim bertanggung jawab atas implementasi", "Laravel jika governance teknis kuat"),
    ("Biaya awal", "Lebih rendah untuk kebutuhan umum", "Lebih tinggi karena pengembangan khusus", "WordPress"),
    ("Biaya jangka panjang", "Dapat naik karena plugin premium/konflik/technical debt", "Prediktif jika codebase dan DevOps dirawat", "Kontekstual"),
    ("Kepemilikan data", "Database mudah, tetapi struktur plugin bervariasi", "Skema eksplisit dan sesuai domain ULT", "Laravel"),
    ("Kemudahan nonteknis", "Sangat baik untuk editor umum", "Filament mudah, tetapi fitur baru tetap butuh developer", "WordPress untuk editorial"),
]
t = doc.add_table(rows=1, cols=4)
table_geometry(t, [1850, 2450, 3200, 1860])
for i, h in enumerate(["Kategori", "WordPress", "Laravel + Filament", "Kecenderungan"]):
    shade(t.cell(0,i), NAVY); font(t.cell(0,i).paragraphs[0].add_run(h), bold=True, color=WHITE)
for row in rows:
    cells = t.add_row().cells
    for i, text in enumerate(row):
        if len(t.rows) % 2 == 0: shade(cells[i], "FBF9FC")
        font(cells[i].paragraphs[0].add_run(text), size=9.2, bold=(i==0), color=NAVY if i==0 else INK)
table_geometry(t, [1850, 2450, 3200, 1860])

page_break()
doc.add_heading("2. Tampilan website dan pengalaman frontend", level=1)
platform_block("WordPress", [
    "Tema dan page builder memungkinkan perubahan layout tanpa menulis banyak kode.",
    "Ekosistem template besar mempersingkat pembuatan landing page, galeri, berita, dan formulir.",
    "Editor dapat membuat halaman kampanye secara mandiri jika blok dan komponen sudah disiapkan.",
], [
    "Kombinasi tema, plugin, dan page builder dapat menghasilkan CSS/JavaScript berlebih serta desain yang tidak konsisten.",
    "Kualitas responsif, aksesibilitas, dan performa sangat bergantung pada tema/plugin yang dipilih.",
    "Perubahan tema atau page builder berpotensi menimbulkan lock-in dan merusak layout lama.",
])
platform_block("Laravel + Filament", [
    "Frontend dapat mengikuti identitas visual ULT secara presisi: Poppins, logo, palet warna, komponen layanan, bilingual, dan aksesibilitas.",
    "Markup, asset loading, SEO metadata, dan perilaku interaktif dapat dikendalikan sampai tingkat kode.",
    "Komponen yang konsisten memudahkan quality assurance lintas halaman dan perangkat.",
], [
    "Perubahan layout/komponen baru umumnya membutuhkan developer dan proses build/deploy.",
    "Tidak ada page builder bebas secara bawaan; fleksibilitas editor sengaja dibatasi agar desain tetap konsisten.",
    "Jika design system tidak terdokumentasi, kode frontend khusus dapat berkembang tidak seragam.",
])
doc.add_heading("Dari POV pengguna umum/mahasiswa", level=2)
bullet("WordPress dapat terasa sangat cepat dan familiar bila tema ringan; sebaliknya dapat terasa berat jika banyak plugin, banner, dan script pihak ketiga.")
bullet("Laravel ULT dapat memberikan alur yang lebih fokus: cari kebutuhan, baca syarat/prosedur, lalu menuju kanal resmi. Pengalaman ini lebih mudah disesuaikan dengan perjalanan mahasiswa.")
bullet("Pada kedua platform, kualitas konten, ketepatan tautan, bahasa yang jelas, dan aksesibilitas lebih menentukan kepuasan pengguna daripada nama framework.")

doc.add_heading("Gap frontend yang masih dapat dikembangkan", level=2)
bullet("Pengujian aksesibilitas formal WCAG 2.2 AA dengan pengguna disabilitas dan audit screen reader.")
bullet("Design system terdokumentasi, komponen Storybook/style guide, dan acceptance criteria responsif.")
bullet("Optimasi gambar otomatis (WebP/AVIF, responsive srcset, CDN) dan Core Web Vitals production.")
bullet("PWA/offline fallback, notifikasi layanan, personalisasi ringan, dan pencarian dengan typo tolerance/sinonim.")

page_break()
doc.add_heading("3. CRUD, pengelolaan konten, dan POV admin", level=1)
platform_block("WordPress", [
    "Post/page/media library, revision, autosave, preview, scheduling, dan editor blok tersedia secara matang.",
    "Banyak admin sudah mengenal dashboard WordPress; onboarding editorial relatif singkat.",
    "Plugin dapat menambah custom post type, form, SEO, workflow, multilingual, dan role dengan cepat.",
], [
    "Semakin banyak plugin CRUD/custom fields, semakin sulit memahami sumber data dan dependensinya.",
    "Hak akses plugin tidak selalu granular atau konsisten; konflik update dapat memengaruhi admin.",
    "Konten layanan yang sangat terstruktur sering dipaksa ke post meta sehingga validasi dan pelaporan menjadi rumit.",
])
platform_block("Laravel + Filament", [
    "Form admin dapat mengikuti struktur layanan ULT: kategori, syarat, dokumen, prosedur, CTA, jadwal, biaya, owner, SEO, status terbit, serta versi ID/EN.",
    "Validasi, filter, bulk action, relasi, dan batasan hapus dapat ditetapkan sesuai aturan bisnis.",
    "Perubahan admin langsung bersumber dari database yang sama dan tampil pada frontend tanpa sinkronisasi manual.",
], [
    "Fitur editorial matang seperti revision diff, approval bertingkat, kalender konten, dan restore mandiri belum otomatis tersedia.",
    "Perubahan struktur form atau business rule membutuhkan perubahan kode, migration, test, dan deployment.",
    "Filament harus dijaga kompatibilitasnya dengan Livewire/Laravel; update paket perlu regression test UI.",
])
doc.add_heading("Yang sudah tersedia di aplikasi Laravel ULT", level=2)
bullet("CRUD tambah, lihat, edit, hapus untuk layanan, kategori, artikel, FAQ, kontak, dan tautan cepat.")
bullet("Konten bilingual Indonesia/English, draft/publish, featured content, metadata SEO, gambar, dan pengurutan.")
bullet("Dark/light mode admin serta dashboard trafik dan status konten.")
bullet("Automated test untuk akses admin, CRUD Livewire, bilingual, security, dan pencatatan trafik.")

doc.add_heading("Gap admin/CMS yang belum terealisasi penuh", level=2)
bullet("Role terpisah Super Admin, Editor, Reviewer, dan Viewer dengan permission granular per resource.")
bullet("Workflow maker-checker: draft -> review -> approved -> scheduled -> published, termasuk komentar reviewer.")
bullet("Version history setiap perubahan, diff, rollback satu klik, dan audit log siapa mengubah apa/kapan.")
bullet("Media library terpusat dengan alt text wajib, crop/resize, kompresi, hak pakai, dan masa berlaku aset.")
bullet("Dashboard kualitas konten: konten kedaluwarsa, tautan rusak, terjemahan belum lengkap, owner kosong, dan halaman jarang dikunjungi.")
bullet("Impersonation yang diaudit, reset password/self-service, 2FA, dan manajemen sesi aktif.")

page_break()
doc.add_heading("4. Backend, arsitektur, API, dan integrasi", level=1)
platform_block("WordPress", [
    "REST API, hooks, cron, user management, media, taxonomy, dan editorial tersedia dari core.",
    "Plugin untuk analytics, form, email, SSO, cache, SEO, backup, dan security sangat banyak.",
    "Cocok untuk tim yang mengutamakan konfigurasi daripada membangun sistem khusus.",
], [
    "Arsitektur dapat menjadi sulit diprediksi ketika banyak plugin mengubah lifecycle request dan database.",
    "Kualitas kode serta pemeliharaan plugin bervariasi; plugin yang ditinggalkan menjadi risiko.",
    "Integrasi kompleks sering tersebar di plugin, functions.php, webhook, dan konfigurasi dashboard.",
])
platform_block("Laravel + Filament", [
    "Separation of concerns, routing, middleware, validation, ORM, queue, event, cache, scheduler, dan testing mendukung aplikasi domain-spesifik.",
    "Integrasi PAuS/SSO, helpdesk, API kampus, notifikasi, atau data warehouse dapat dibuat sebagai service yang eksplisit dan teruji.",
    "Struktur kode, migration, dan test dapat direview serta ditelusuri melalui version control.",
], [
    "Semua integrasi khusus menjadi tanggung jawab tim: desain, failure handling, retry, observability, dan dokumentasi.",
    "Kualitas arsitektur dapat turun bila logic ditempatkan sembarang atau tidak ada code review.",
    "Queue worker, scheduler, cache store, dan object storage membutuhkan pengelolaan infrastruktur tambahan saat digunakan.",
])
doc.add_heading("Integrasi lanjutan yang potensial", level=2)
bullet("SSO Unpad/PAuS untuk admin dan, bila diperlukan, pengalaman mahasiswa terautentikasi.")
bullet("Sinkronisasi status atau deep-link terkontrol ke helpdesk.unpad.ac.id tanpa menjadikan portal ini sistem ticketing duplikat.")
bullet("API katalog layanan untuk aplikasi Unpad lain, chatbot resmi, atau kanal mobile.")
bullet("Webhook/notifikasi saat konten kritikal berubah, link gagal, atau masa review konten berakhir.")
bullet("OpenTelemetry/APM, centralized logging, error tracking, dan uptime monitoring production.")

page_break()
doc.add_heading("5. Database, keamanan, dan privasi", level=1)
doc.add_heading("Database", level=2)
platform_block("WordPress", [
    "Skema core stabil, terdokumentasi, dan mudah didukung banyak hosting.",
    "Backup/migration tooling tersedia luas; editorial post/revision sudah terintegrasi.",
], [
    "Plugin banyak memakai wp_options dan post meta; data terstruktur dapat sulit divalidasi, di-query, dan dibersihkan.",
    "Uninstall plugin tidak selalu menghapus tabel/options; database dapat menumpuk data yatim.",
    "Relasi dan constraint domain layanan tidak sejelas skema relasional khusus.",
])
platform_block("Laravel + Filament", [
    "Skema tabel ULT eksplisit, migration terlacak, foreign key dapat ditegakkan, dan query domain lebih jelas.",
    "Model layanan, kategori, artikel, FAQ, kontak, analytics, dan user dapat diindeks sesuai pola akses.",
    "Pemisahan database production/staging/test serta backup terjadwal dapat dirancang dengan jelas.",
], [
    "Migration yang salah dapat merusak data jika tidak diuji, dibackup, dan memiliki rollback plan.",
    "Developer harus menghindari N+1 query, indeks yang kurang, retention tanpa batas, dan penyimpanan data sensitif berlebihan.",
    "Keamanan database bergantung pada kredensial, network policy, encryption, backup, dan hak akses server; bukan ORM saja.",
])

doc.add_heading("Security", level=2)
platform_block("WordPress", [
    "Core aktif dipelihara, patch keamanan tersedia luas, dan hardening guidance matang.",
    "Security plugin/WAF dapat memberi proteksi cepat seperti brute-force limiting, scanning, dan 2FA.",
], [
    "Popularitas WordPress menjadikannya target otomatis; plugin/theme rentan atau tidak diperbarui adalah jalur serangan umum.",
    "Admin plugin/theme editor, XML-RPC, akun lama, shared hosting, dan permission file yang buruk memperbesar risiko.",
    "Supply-chain plugin premium/bajakan dan kredensial admin lemah adalah risiko operasional signifikan.",
])
platform_block("Laravel + Filament", [
    "CSRF, session security, password hashing, validation, ORM parameter binding, policy/gate, dan middleware tersedia sebagai fondasi.",
    "Permukaan aplikasi dapat dibatasi hanya pada fitur yang dipakai; tidak perlu membuka ekosistem plugin publik.",
    "Security header, admin-domain restriction, hash analytics, dan test security sudah menjadi bagian aplikasi ULT saat ini.",
], [
    "Kerentanan business logic, authorization, upload file, XSS dari rich text, SSRF, secret leakage, dan dependency tetap mungkin terjadi.",
    "Tidak ada patch otomatis yang cukup tanpa proses composer audit, dependency update, code review, pentest, dan monitoring.",
    "Panel admin custom harus diuji setelah upgrade Laravel/Filament/Livewire karena regresi frontend dapat menghambat operasi.",
])

callout("Prinsip penting", "Laravel tidak otomatis lebih aman daripada WordPress, dan WordPress tidak otomatis tidak aman. Laravel memberi kontrol lebih besar dan permukaan yang lebih spesifik; kontrol tersebut hanya menjadi keunggulan bila ada governance teknis, patching, backup, monitoring, least privilege, dan pengujian berkala.", fill="FFF4E5", accent=GOLD)

doc.add_heading("Kontrol keamanan yang masih perlu dituntaskan sebelum produksi", level=2)
for x in [
    "2FA admin, rate limiting login, lockout adaptif, session timeout, dan revocation sesi.",
    "Role/permission granular dan audit log immutable untuk perubahan konten serta konfigurasi.",
    "Validasi upload MIME/content, antivirus scanning, image re-encoding, private object storage, dan signed URL bila diperlukan.",
    "Content Security Policy yang diuji, HSTS production, secure cookie, trusted proxy/host, dan secret rotation.",
    "Database private network, akun DB least privilege, encryption at rest/in transit, backup terenkripsi, serta restore drill.",
    "SAST/dependency scan/secret scan di CI, DAST pada staging, penetration test sebelum go-live, dan incident response runbook.",
    "Kebijakan retention analytics, privacy notice, mekanisme penghapusan, serta kajian kepatuhan data institusi.",
]: bullet(x)

page_break()
doc.add_heading("6. SEO, performa, aksesibilitas, dan analytics", level=1)
doc.add_heading("SEO", level=2)
bullet("WordPress unggul lewat plugin SEO siap pakai, sitemap, schema, redirect manager, dan preview editorial. Risiko muncul bila beberapa plugin SEO aktif bersamaan.")
bullet("Laravel memberi kontrol penuh atas metadata, canonical, sitemap, structured data, redirect, dan cache, tetapi fitur harus dibuat/dikonfigurasi serta diuji sendiri.")
bullet("Gap Laravel ULT: sitemap otomatis, robots production, canonical audit, redirect map WordPress -> Laravel, structured data Organization/Article/FAQ, dan Search Console verification.")

doc.add_heading("Performa", level=2)
bullet("WordPress dapat cepat dengan full-page cache, CDN, image optimization, tema ringan, dan plugin minimal; dapat sangat lambat bila query/plugin/script menumpuk.")
bullet("Laravel dapat mengoptimalkan query, config/route/view cache, Redis, queue, CDN, dan eager loading; tetapi salah konfigurasi atau dashboard query berat juga dapat memperlambat.")
bullet("Gap Laravel ULT: baseline Lighthouse production, load test, CDN, Redis/cache strategy, image pipeline, database slow-query monitoring, dan performance budget CI.")

doc.add_heading("Aksesibilitas", level=2)
bullet("WordPress memiliki tema/plugin aksesibilitas, tetapi widget overlay tidak menggantikan semantic HTML dan remediation nyata.")
bullet("Laravel ULT sudah memiliki menu aksesibilitas custom dan semantic controls; kelebihannya dapat terintegrasi dengan desain, namun tetap memerlukan audit manual dan pengguna nyata.")
bullet("Gap: VPAT/accessibility statement, keyboard/screen-reader test matrix, caption/transcript media, alt-text governance, dan proses pengaduan aksesibilitas.")

doc.add_heading("Analytics", level=2)
bullet("WordPress mudah dihubungkan ke GA4/Matomo/plugin analytics, tetapi dapat menambah script pihak ketiga dan implikasi consent/privacy.")
bullet("Laravel ULT sudah mencatat page view, pengunjung unik berbasis hash, perangkat, referrer, pencarian, dan klik keluar tanpa menyimpan IP mentah.")
bullet("Gap: retention/aggregation job, filter internal staff, export laporan, goal conversion, consent decision, data dictionary, serta rekonsiliasi dengan log reverse proxy/CDN.")

page_break()
doc.add_heading("7. Operasional, deployment, SDM, dan biaya", level=1)
platform_block("WordPress", [
    "Banyak pilihan managed hosting dan operator yang familiar.",
    "Update konten tidak memerlukan deployment kode; plugin/theme dapat diperbarui lewat dashboard.",
    "Biaya awal rendah untuk kebutuhan situs informasi standar.",
], [
    "Update plugin langsung di production dapat menimbulkan konflik; tetap membutuhkan staging, backup, dan rollback.",
    "Biaya plugin premium, lisensi, support, dan technical debt dapat bertambah tersembunyi.",
    "Kepemilikan konfigurasi sering tersebar di database/dashboard sehingga perubahan sulit direview sebagai kode.",
])
platform_block("Laravel + Filament", [
    "Deployment dapat direproduksi melalui Git, Composer lock, migration, build assets, CI/CD, dan environment configuration.",
    "Testing serta code review memberi jejak perubahan yang kuat dan mendukung rollback release.",
    "Lebih cocok untuk roadmap aplikasi dan integrasi institusional jangka panjang.",
], [
    "Memerlukan kompetensi PHP/Laravel, database, Linux/web server, CI/CD, security, dan observability.",
    "Hosting harus mendukung worker/scheduler/storage/cache sesuai fitur; shared hosting minimal dapat membatasi.",
    "Biaya awal dan tanggung jawab pemeliharaan lebih tinggi daripada website template sederhana.",
])

doc.add_heading("Kebutuhan minimum operasi Laravel ULT", level=2)
numbered("Environment staging terpisah dari production dengan data tersanitasi.")
numbered("CI menjalankan test, formatting, dependency/security scan, dan build frontend.")
numbered("Deployment atomic atau maintenance window, migration plan, health check, dan rollback teruji.")
numbered("Database managed/private, backup harian, object storage backup, retention, dan restore drill.")
numbered("Monitoring uptime, error, latency, queue, storage, certificate, dan database; alert memiliki PIC/escalation path.")
numbered("Jadwal patch bulanan dan emergency patch untuk vulnerability kritikal.")

page_break()
doc.add_heading("8. Migrasi WordPress ke Laravel pada domain yang sama", level=1)
doc.add_paragraph("Pergantian platform tidak cukup dengan mengganti file website. Migrasi harus menjaga URL, konten, SEO, analytics, media, email/form, akses admin, dan kemampuan rollback.")
doc.add_heading("Risiko utama", level=2)
bullet("URL lama berubah dan menghasilkan 404 sehingga trafik serta ranking turun.")
bullet("Artikel/media/metadata tidak terimpor lengkap atau encoding rusak.")
bullet("Form, email, webhook, analytics, sitemap, dan redirect lama terlewat.")
bullet("DNS, TLS, cache CDN, session, permission storage, queue, atau migration database salah saat cutover.")
bullet("Tidak ada snapshot WordPress sehingga rollback sulit.")

doc.add_heading("Checklist cutover", level=2)
for x in [
    "Inventaris seluruh URL WordPress, status HTTP, title/meta, canonical, media, dan trafiknya.",
    "Buat mapping konten dan redirect 301 satu-per-satu; jangan mengarahkan seluruh 404 ke homepage.",
    "Lakukan dry-run import dan validasi jumlah record, checksum media, bahasa, author, tanggal, serta internal link.",
    "Freeze editorial singkat, backup penuh database/files WordPress, dan tetapkan rollback window.",
    "Deploy Laravel di origin/server baru, uji menggunakan staging domain/hosts override, lalu alihkan reverse proxy/DNS.",
    "Pantau 4xx/5xx, latency, indexing, form, login admin, storage, email, queue, dan conversion setelah cutover.",
    "Pertahankan WordPress lama dalam mode private/read-only selama periode rollback yang disepakati, lalu arsipkan secara aman.",
]: bullet(x)

page_break()
doc.add_heading("9. Gap analysis aplikasi Laravel ULT saat ini", level=1)
gaps = [
    ("Prioritas 0 - sebelum go-live", "Security & reliability", "2FA, role/permission, audit log, backup+restore drill, production hardening, monitoring, pentest, redirect map, staging/CI/CD", "Risiko akses tidak sah, kehilangan data, downtime, dan penurunan SEO"),
    ("Prioritas 1", "Editorial governance", "Approval workflow, revision/rollback, content expiry, owner/reviewer, link checker, translation completeness", "Kualitas dan akurasi informasi lebih terjaga"),
    ("Prioritas 1", "Performance & media", "CDN, WebP/AVIF, image resize, caching, Redis, query monitoring, load test", "Waktu muat dan kapasitas lebih baik"),
    ("Prioritas 1", "Accessibility", "Audit WCAG 2.2 AA, screen reader, keyboard, user testing, accessibility statement", "Layanan lebih inklusif dan dapat dipertanggungjawabkan"),
    ("Prioritas 2", "Analytics maturity", "Retention, aggregation, export, conversion funnel, internal traffic filter, dashboard date filter", "Keputusan konten berbasis data"),
    ("Prioritas 2", "Integrasi", "SSO PAuS, API katalog, helpdesk deep-link/status, notification/webhook", "Mengurangi duplikasi dan friksi layanan"),
    ("Prioritas 2", "Admin productivity", "Media library, bulk translation, scheduled review reminders, import/export", "Operasi konten lebih cepat"),
    ("Prioritas 3", "Experience innovation", "PWA, rekomendasi layanan, chatbot terkontrol, saved services, notification preference", "Pengalaman mahasiswa lebih personal"),
]
t = doc.add_table(rows=1, cols=4)
table_geometry(t, [1700, 1700, 3900, 2060])
for i,h in enumerate(["Prioritas", "Area", "Belum terealisasi / perlu dituntaskan", "Dampak"]):
    shade(t.cell(0,i), NAVY); font(t.cell(0,i).paragraphs[0].add_run(h), bold=True, color=WHITE, size=9.3)
for row in gaps:
    c=t.add_row().cells
    for i,text in enumerate(row):
        if len(t.rows)%2==0: shade(c[i], "FBF9FC")
        font(c[i].paragraphs[0].add_run(text), size=9, bold=(i<2), color=NAVY if i<2 else INK)
table_geometry(t, [1700, 1700, 3900, 2060])

doc.add_heading("Definisi selesai yang disarankan", level=2)
bullet("Fitur dianggap selesai bila memiliki acceptance criteria, authorization, validation, auditability, test otomatis, black-box test, dokumentasi operator, monitoring, dan rollback plan.")
bullet("Go-live dianggap siap bila checklist security, data migration, SEO, accessibility, performance, backup/restore, dan incident response ditandatangani PIC terkait.")

page_break()
doc.add_heading("10. Rekomendasi keputusan", level=1)
callout("Rekomendasi", "Gunakan Laravel + Filament sebagai platform utama ULT dan pertahankan WordPress lama hanya sebagai sumber migrasi/arsip sementara. Keputusan ini tepat bila Unpad menetapkan pemilik teknis, pemilik konten, anggaran operasi, dan standar keamanan yang jelas.")
doc.add_heading("Alasan", level=2)
bullet("Model layanan ULT lebih terstruktur daripada pola website berita biasa.")
bullet("Bilingual, aksesibilitas, dashboard trafik, CRUD domain, dan integrasi membutuhkan kontrol aplikasi yang kuat.")
bullet("Skema data dan business rule eksplisit lebih mudah diuji, diaudit, dan dikembangkan.")
bullet("Arah jangka panjang berpotensi melibatkan SSO, helpdesk, API, workflow, dan analytics yang lebih kompleks.")

doc.add_heading("Kapan WordPress tetap lebih tepat", level=2)
bullet("Jika kebutuhan hanya website publik sederhana dengan berita/halaman, perubahan layout sering oleh nondeveloper, dan tidak ada integrasi/domain logic khusus.")
bullet("Jika organisasi belum memiliki kapasitas untuk memelihara aplikasi Laravel, CI/CD, database, monitoring, dan security patching.")
bullet("Jika target waktu/anggaran sangat terbatas dan risiko plugin dapat dikendalikan melalui managed WordPress serta governance ketat.")

doc.add_heading("Roadmap 90 hari yang disarankan", level=2)
numbered("Hari 1-30: audit WordPress dan URL, hardening Laravel, role/2FA/audit log, staging, backup/restore, monitoring, dan baseline accessibility/performance.")
numbered("Hari 31-60: dry-run migrasi, redirect map, revision/approval workflow minimum, media optimization, load/security test, dan pelatihan admin.")
numbered("Hari 61-90: UAT stakeholder, cutover rehearsal, go-live dengan rollback window, observasi intensif, perbaikan 4xx/5xx, dan review KPI awal.")

doc.add_heading("Batasan kajian", level=1)
doc.add_paragraph("Dokumen ini bukan hasil penetration test, audit source code WordPress lama, audit infrastruktur production, legal opinion, atau pengukuran biaya vendor. Sebelum keputusan final, lakukan inventaris plugin/tema/versi WordPress, scan kerentanan, review hosting, pengukuran trafik, pemeriksaan database, serta validasi kebutuhan stakeholder ULT, Direktorat TIK, keamanan informasi, komunikasi, dan layanan disabilitas.")

# Keep rows together where practical and repeat table headers
for table in doc.tables:
    if table.rows:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        hdr = OxmlElement("w:tblHeader"); hdr.set(qn("w:val"), "true"); trPr.append(hdr)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit"); trPr.append(cant)

doc.core_properties.title = "Perbandingan WordPress vs Laravel + Filament untuk Website ULT Unpad"
doc.core_properties.subject = "Kajian platform, risiko, keamanan, dan roadmap pengembangan"
doc.core_properties.author = "Unit Layanan Terpadu Universitas Padjadjaran"
doc.core_properties.keywords = "ULT Unpad, WordPress, Laravel, Filament, CMS, security, database"
doc.save(OUT)
print(OUT)
