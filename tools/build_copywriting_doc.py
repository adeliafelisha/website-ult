from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "ULT_Unpad_Inventaris_Konten_Website.docx"
PURPLE = RGBColor(112, 29, 107)
DARK = RGBColor(54, 18, 51)
GOLD = RGBColor(213, 157, 22)
MUTED = RGBColor(95, 85, 94)

def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def keep(p):
    p.paragraph_format.keep_with_next = True

def label(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    font(p.add_run(text.upper()), 8.5, True, GOLD)
    return p

def h1(doc, text): return doc.add_heading(text, level=1)
def h2(doc, text): return doc.add_heading(text, level=2)
def h3(doc, text): return doc.add_heading(text, level=3)
def para(doc, text, bold_lead=None):
    p=doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), 11, True, DARK); font(p.add_run(text[len(bold_lead):]), 11)
    else: font(p.add_run(text), 11)
    return p
def bullet(doc, text):
    p=doc.add_paragraph(style="List Bullet"); font(p.add_run(text),11); return p
def numbered(doc, text):
    p=doc.add_paragraph(style="List Number"); font(p.add_run(text),11); return p
def field(doc, name, value):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    font(p.add_run(name+": "),10.5,True,DARK); font(p.add_run(value or "—"),10.5)
def page_break(doc): doc.add_page_break()

doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Heading 1",16,PURPLE,18,10),("Heading 2",13,PURPLE,14,7),("Heading 3",12,DARK,10,5)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for lname in ["List Bullet","List Number"]:
    s=styles[lname]; s.font.name="Calibri"; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.375); s.paragraph_format.first_line_indent=Inches(-.188); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25

# Header and footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(hp.add_run("ULT UNPAD  |  INVENTARIS KONTEN"),8.5,True,MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(fp.add_run("Dokumen pemeriksaan copywriting • 12 Agustus 2026"),8,color=MUTED)

# Cover: editorial_cover pattern with ULT palette override
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(120); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("INVENTARIS KONTEN WEBSITE"),10,True,GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); font(p.add_run("Unit Layanan Terpadu Unpad"),30,True,PURPLE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(30); font(p.add_run("Salinan lengkap untuk pemeriksaan copywriting manual"),14,color=DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Cakupan: navigasi global, Beranda, Profil ULT, Layanan, Artikel, FAQ, Kontak, Pencarian, aksesibilitas, footer, serta versi Bahasa Inggris."),10.5,color=MUTED)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(84); font(p.add_run("STATUS: DRAF KONTEN APLIKASI"),9,True,GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Disusun dari source code dan data awal aplikasi Laravel"),9,color=MUTED)
page_break(doc)

h1(doc,"Petunjuk Pemeriksaan")
para(doc,"Dokumen ini menyalin isi konten website per menu dan mengikuti urutan baca pengguna: elemen global, isi halaman, kondisi kosong, lalu footer. Teks dinamis dari CMS dicatat berdasarkan data awal aplikasi. Perubahan yang dibuat admin setelah dokumen ini diterbitkan tidak otomatis masuk ke dokumen.")
label(doc,"Penanda")
bullet(doc,"[DINAMIS/CMS] — teks dapat berubah melalui Filament Admin.")
bullet(doc,"[CONTOH/VERIFIKASI] — data belum boleh dianggap resmi sebelum diverifikasi.")
bullet(doc,"[KONDISI KOSONG] — hanya tampil saat data belum tersedia atau pencarian tidak menemukan hasil.")
bullet(doc,"[EN] — salinan Bahasa Inggris yang tersedia pada antarmuka.")
h2(doc,"Daftar Menu")
for x in ["Elemen Global","Beranda","Profil ULT","Layanan dan Detail Layanan","Artikel dan Detail Artikel","FAQ","Kontak","Pencarian","Versi Bahasa Inggris"]: bullet(doc,x)

page_break(doc); h1(doc,"1. Elemen Global")
h2(doc,"1.1 Header dan Navigasi")
for x in ["Lewati ke konten utama","ULT Unpad","Unit Layanan Terpadu","Beranda","Profil ULT","Layanan","Artikel","FAQ","Kontak","Buka pencarian","Aksesibilitas","Bahasa: ID / EN","Buka menu"]: bullet(doc,x)
h2(doc,"1.2 Panel Pencarian Global")
field(doc,"Pertanyaan","Apa yang ingin Anda cari?"); field(doc,"Placeholder","Contoh: KTM hilang, UKT, ijazah..."); field(doc,"Tombol","Cari")
h2(doc,"1.3 Menu Aksesibilitas")
field(doc,"Judul","Menu Aksesibilitas (CTRL+U)"); para(doc,"Sesuaikan tampilan website dengan kebutuhan Anda. Pengaturan tersimpan otomatis.")
for x in ["Perbesar Teks","Perkecil Teks","Kontras Tinggi","Mode Gelap","Sorot Tautan","Spasi Teks","Hentikan Animasi","Sembunyikan Gambar","Ramah Disleksia","Kursor Besar","Tinggi Baris","Perataan Teks","Kejenuhan","Atur Ulang Semua Pengaturan","Tutup menu aksesibilitas"]: bullet(doc,x)
h2(doc,"1.4 Footer")
field(doc,"Merek","ULT Unpad — Melayani dengan PASTI"); para(doc,"Portal informasi dan pengarah layanan untuk sivitas akademika dan masyarakat.")
field(doc,"Jelajahi","Profil ULT; Direktori Layanan; Artikel & Informasi; Pertanyaan Umum")
field(doc,"Tautan Utama","Website Unpad; SMUP; Hubungi ULT")
field(doc,"Lokasi","Gedung Rektorat Unpad, Jatinangor, Sumedang")
para(doc,"Informasi kontak perlu diverifikasi admin sebelum produksi.")
field(doc,"Bagian bawah","© 2026 ULT Universitas Padjadjaran; Masuk Admin")

page_break(doc); h1(doc,"2. Beranda")
h2(doc,"2.1 Hero")
label(doc,"Badge"); para(doc,"Satu pintu informasi layanan Unpad")
label(doc,"Judul utama"); para(doc,"Temukan layanan kampus tanpa bingung.")
para(doc,"Mulai dari kebutuhan akademik, mahasiswa baru, teknologi informasi, hingga layanan inklusif—kami arahkan Anda ke langkah yang tepat.")
field(doc,"Placeholder pencarian","Cari layanan, informasi, atau pertanyaan..."); field(doc,"Tombol","Cari sekarang"); field(doc,"Pencarian populer","KTM; UKT; Ijazah")
field(doc,"Kartu gambar","Yang paling sering dibutuhkan mahasiswa — Daring, luring, dan hybrid")
h2(doc,"2.2 Kategori Layanan")
label(doc,"Pengantar"); para(doc,"Mulai dari kebutuhan Anda"); para(doc,"Saya ingin mencari layanan untuk..."); field(doc,"Tautan","Lihat semua layanan")
cats=[("Akademik Mahasiswa","Registrasi, UKT, KRS, dan dokumen mahasiswa."),("Mahasiswa Baru & SMUP","Penerimaan, verifikasi, NPM, dan KTM."),("Dosen & Tendik","Layanan dosen dan tenaga kependidikan."),("Layanan Internasional","Exchange, visa, KITAS, dan study permit."),("PAuS & Teknologi Informasi","Akun PAuS, SIAT, PINTAS, dan LiVE."),("Layanan Disabilitas","Akses fasilitas dan pendampingan inklusif.")]
for n,d in cats: h3(doc,n); para(doc,d); para(doc,"[DINAMIS/CMS] {jumlah} layanan tersedia →")
para(doc,"[KONDISI KOSONG] Data kategori akan tampil setelah admin menambahkannya.")
h2(doc,"2.3 Layanan Utama")
label(doc,"Pengantar"); para(doc,"Layanan utama"); para(doc,"Yang paling sering dibutuhkan mahasiswa"); field(doc,"Tautan","Jelajahi semuanya")
para(doc,"[DINAMIS/CMS] Kartu berisi jenis layanan, judul, ringkasan, sasaran pengguna, dan ajakan “Lihat detail”.")
para(doc,"[KONDISI KOSONG] Konten layanan sedang disiapkan oleh admin.")
h2(doc,"2.4 Kampus Inklusif")
label(doc,"Pengantar"); para(doc,"Kampus inklusif"); para(doc,"Layanan yang dapat diakses oleh semua")
para(doc,"ULT Unpad membantu pengguna dengan disabilitas menemukan informasi fasilitas, alternatif pelayanan daring, dan kanal pendampingan yang sesuai.")
for x in ["Informasi akses lokasi dan fasilitas","Alternatif kanal layanan daring","Konten ramah pembaca layar dan keyboard"]: bullet(doc,x)
field(doc,"Tombol","Lihat layanan disabilitas")
h2(doc,"2.5 Artikel Terbaru")
label(doc,"Pengantar"); para(doc,"Artikel terbaru"); para(doc,"Informasi penting untuk kehidupan kampus"); field(doc,"Tautan","Lihat semua artikel")
para(doc,"[DINAMIS/CMS] Kartu artikel menampilkan kategori, judul, ringkasan, tanggal terbit, dan penulis.")
para(doc,"[KONDISI KOSONG] Artikel akan tampil setelah diterbitkan admin.")
h2(doc,"2.6 FAQ Ringkas")
label(doc,"Pengantar"); para(doc,"Jawaban cepat"); para(doc,"Pertanyaan yang sering diajukan")
para(doc,"Tidak menemukan jawaban? Gunakan pencarian atau hubungi kanal resmi sesuai kebutuhan Anda."); field(doc,"Tombol","Baca semua FAQ"); para(doc,"[KONDISI KOSONG] FAQ sedang disiapkan.")
h2(doc,"2.7 Portal dan Kanal Resmi")
label(doc,"Pengantar"); para(doc,"Lanjutkan ke sistem resmi"); para(doc,"Portal & kanal layanan Unpad")
for n,d,u in [("Website Unpad","Informasi utama Universitas Padjadjaran","https://www.unpad.ac.id"),("SMUP","Seleksi masuk dan registrasi","https://smup.unpad.ac.id"),("PAuS","Akses akun dan sistem terintegrasi","https://paus.unpad.ac.id"),("PPID Unpad","Informasi publik","https://ppid.unpad.ac.id")]: h3(doc,n); para(doc,d); field(doc,"URL",u)

page_break(doc); h1(doc,"3. Profil ULT")
h2(doc,"3.1 Hero Profil")
label(doc,"Pengantar"); para(doc,"Tentang ULT Unpad"); para(doc,"Garda terdepan pelayanan Unpad.")
para(doc,"Sejak 2015, ULT Unpad terus berkembang menjadi pintu pelayanan publik yang terpadu, mudah diakses, dan berorientasi pada kebutuhan pengguna.")
field(doc,"Sorotan","2015 — Mulai berdiri; 2016 — Layanan satu pintu; PASTI — Nilai pelayanan")
field(doc,"Navigasi bagian","Sejarah; Visi & Misi; PASTI; Dasar Hukum; Petugas")
h2(doc,"3.2 Sejarah")
label(doc,"Pengantar"); para(doc,"Perjalanan kami"); para(doc,"Tumbuh dari satu atap menjadi satu pintu")
para(doc,"Unit Layanan Terpadu (ULT) Unpad didirikan pada tahun 2015 dengan sistem pelayanan “Satu Atap”. Pada masa awal, layanan dilaksanakan dalam satu ruangan sederhana oleh perwakilan dari tiap unit kerja yang ditugaskan di ULT.")
para(doc,"Selaras dengan rencana strategis Unpad untuk meningkatkan kualitas kelembagaan dan prinsip Peningkatan Mutu Berkelanjutan, ULT terus dikembangkan untuk memotong hambatan pelayanan antarsatuan kerja.")
h3(doc,"2015 — Pelayanan Satu Atap"); para(doc,"Perwakilan berbagai unit kerja melayani pengguna bersama-sama dalam satu lokasi.")
h3(doc,"22 Agustus 2016 — Transformasi Menjadi Satu Pintu"); para(doc,"ULT bertransformasi menjadi layanan “Satu Pintu”, dengan proses pengelolaan permohonan hingga penyelesaian dilakukan dalam satu tempat.")
h3(doc,"Kini — Layanan Hybrid dan Inklusif"); para(doc,"ULT menggabungkan pelayanan luring dan daring serta terus memperkuat pelayanan yang ramah disabilitas.")
h2(doc,"3.3 Visi dan Misi")
label(doc,"Pengantar"); para(doc,"Arah pelayanan"); para(doc,"Visi dan misi")
h3(doc,"Visi"); para(doc,"Terwujudnya pelayanan informasi publik di Universitas Padjadjaran yang profesional, akuntabel, simpel, transparan, dan informatif untuk memenuhi hak pemohon informasi publik sesuai ketentuan peraturan perundang-undangan.")
h3(doc,"Misi")
for x in ["Menyediakan pengelolaan dan pelayanan publik yang cepat, tepat waktu, sederhana, dan berkualitas.","Membangun dan mengembangkan sistem pengelolaan dan layanan informasi yang efektif dan efisien berbasis teknologi informasi.","Meningkatkan kualitas penyelenggaraan pemerintahan melalui optimalisasi pengelolaan dan pelayanan publik satu pintu."]: numbered(doc,x)
h2(doc,"3.4 Motto PASTI")
para(doc,"Profesional, Akuntabel, Simpel, Transparan, Informatif.")
for n,d in [("P — Profesional","Pelayanan kompeten dan responsif."),("A — Akuntabel","Bertanggung jawab dan terukur."),("S — Simpel","Proses yang mudah dipahami."),("T — Transparan","Syarat dan prosedur yang jelas."),("I — Informatif","Informasi yang relevan dan berguna.")]: h3(doc,n); para(doc,d)
h2(doc,"3.5 Dasar Hukum")
label(doc,"Pengantar"); para(doc,"Pelayanan yang akuntabel"); para(doc,"Sebagai komitmen terhadap pendirian dan pemenuhan kepuasan pengguna layanan, Unpad menetapkan landasan hukum untuk pembentukan dan pengelolaan ULT.")
for n,d in [("Peraturan Rektor No. 1 Tahun 2020","Tentang Susunan Organisasi dan Tata Kelola Universitas Padjadjaran."),("Keputusan Rektor No. 1289/UN6,RKT/Kep/HK/2016","Tentang Pembentukan Unit Layanan Terpadu Universitas Padjadjaran."),("Keputusan Rektor No. 2086/UN6,RKT/Kep/HK/2017","Tentang Panduan Pengelolaan Unit Layanan Terpadu Universitas Padjadjaran.")]: h3(doc,n); para(doc,d)
h2(doc,"3.6 Petugas Penyelenggara")
label(doc,"Keterangan foto"); para(doc,"Melayani dengan empati"); label(doc,"Pengantar"); para(doc,"Petugas penyelenggara"); para(doc,"Disiapkan untuk memberi layanan prima")
para(doc,"Petugas Unit Layanan Terpadu Universitas Padjadjaran melalui proses seleksi yang ketat, termasuk asesmen psikotes dan pemagangan pada unit kerja di berbagai direktorat.")
para(doc,"Para petugas dibekali kompetensi layanan prima melalui pelatihan berkala agar ULT dapat menjadi garda terdepan yang andal bagi dosen, tenaga kependidikan, mahasiswa, dan masyarakat.")
field(doc,"Tombol","Hubungi ULT"); field(doc,"Keterangan galeri","Ruang layanan yang nyaman; Ruang publik yang inklusif; Pintu layanan yang jelas")

page_break(doc); h1(doc,"4. Layanan")
h2(doc,"4.1 Halaman Direktori")
label(doc,"Pengantar"); para(doc,"Direktori layanan"); para(doc,"Temukan layanan sesuai kebutuhan Anda")
para(doc,"Gunakan filter kategori dan sasaran pengguna untuk mempersempit hasil.")
field(doc,"Filter kategori","Semua kategori + kategori dari CMS"); field(doc,"Filter sasaran","Semua pengguna; Mahasiswa; Mahasiswa Baru; Dosen & Tendik; Pengguna Internasional; Pengguna dengan Disabilitas"); field(doc,"Tombol","Terapkan filter")
para(doc,"[KONDISI KOSONG] Layanan belum ditemukan — Coba ubah filter atau hubungi ULT untuk diarahkan ke kanal yang tepat.")
h2(doc,"4.2 Pola Detail Layanan")
for x in ["Breadcrumb: Beranda / Layanan / {judul}","Jenis layanan: Daring, Luring, atau Hybrid","Judul dan ringkasan","Tombol tindakan","Siapa yang dapat menggunakan?","Persyaratan","Dokumen yang diperlukan","Prosedur layanan","Informasi layanan: jenis, lokasi, waktu layanan, estimasi proses, biaya, unit penanggung jawab","Terakhir diperbarui; pemilik konten","Layanan terkait"]: bullet(doc,x)
para(doc,"Teks default sasaran: Sivitas akademika dan masyarakat sesuai ketentuan layanan.")
para(doc,"Teks default biaya: Tidak dipungut biaya / sesuai ketentuan.")

services=[
("Penggantian KTM Hilang atau Rusak","Akademik Mahasiswa","Hybrid","Panduan persyaratan dan langkah penggantian Kartu Tanda Mahasiswa.","Mahasiswa Aktif","Surat kehilangan dari kepolisian untuk KTM hilang; Identitas mahasiswa aktif.","Surat kehilangan, KTP, dan pas foto terbaru.",["Siapkan seluruh dokumen.","Hubungi ULT untuk verifikasi.","Ikuti arahan pada kanal resmi."],"Hubungi Admin ULT","—"),
("Informasi UKT dan Registrasi","Akademik Mahasiswa","Daring","Informasi pembayaran UKT, heregistrasi, dan kanal resmi bantuan.","Mahasiswa Aktif","Status mahasiswa aktif dan data pembayaran.","Bukti pembayaran jika diperlukan.",["Periksa tagihan pada sistem akademik.","Lakukan pembayaran melalui kanal resmi.","Hubungi ULT bila status belum diperbarui."],"Buka PAuS","https://paus.unpad.ac.id"),
("Registrasi Mahasiswa Baru","Mahasiswa Baru & SMUP","Daring","Panduan registrasi, verifikasi dokumen, dan akses informasi penerimaan.","Mahasiswa Baru","Dinyatakan diterima melalui jalur resmi.","Dokumen sesuai ketentuan portal SMUP.",["Buka portal SMUP.","Masuk dengan akun peserta.","Ikuti tahapan dan tenggat."],"Buka Website SMUP","https://smup.unpad.ac.id"),
("Pendampingan Layanan Disabilitas","Layanan Disabilitas","Hybrid","Informasi akses fasilitas dan pendampingan layanan kampus.","Pengguna dengan Disabilitas","Sampaikan kebutuhan akses yang diperlukan.","Dokumen pendukung bila relevan.",["Pelajari opsi akses.","Hubungi ULT sebelum kunjungan.","Sepakati pendampingan yang dibutuhkan."],"Hubungi Admin ULT","—"),
("Bantuan Akun PAuS","PAuS & Teknologi Informasi","Daring","Arah layanan kendala akun, lupa kata sandi, dan akses sistem.","Mahasiswa, Dosen & Tendik","Memiliki identitas sivitas Unpad.","NPM/NIP dan identitas pendukung.",["Buka PAuS.","Gunakan fitur pemulihan akun.","Gunakan Support Unpad jika perlu."],"Buka PAuS","https://paus.unpad.ac.id"),
("Informasi Mahasiswa Internasional","Layanan Internasional","Daring","Routing informasi exchange, visa, KITAS, dan study permit.","Pengguna Internasional","Sesuai program dan status studi.","Paspor dan dokumen program.",["Periksa informasi program.","Siapkan dokumen perjalanan.","Hubungi International Office."],"Buka International Office","https://international.unpad.ac.id")]
for s in services:
    h2(doc,"4.3 "+s[0]); field(doc,"Kategori",s[1]); field(doc,"Jenis",s[2]); field(doc,"Ringkasan",s[3]); field(doc,"Sasaran",s[4]); field(doc,"Persyaratan",s[5]); field(doc,"Dokumen",s[6]); h3(doc,"Prosedur"); [numbered(doc,x) for x in s[7]]; field(doc,"CTA",s[8]); field(doc,"URL",s[9]); field(doc,"Pemilik konten","ULT Unpad"); field(doc,"Unit penanggung jawab","Unit terkait — perlu verifikasi")

page_break(doc); h1(doc,"5. Artikel")
h2(doc,"5.1 Halaman Daftar Artikel")
label(doc,"Pengantar"); para(doc,"Artikel & informasi"); para(doc,"Panduan ringkas untuk kehidupan kampus")
para(doc,"Informasi layanan, panduan mahasiswa, aksesibilitas, dan kabar terbaru ULT Unpad.")
para(doc,"[DINAMIS/CMS] Kartu menampilkan kategori, judul, ringkasan, tanggal terbit, dan penulis.")
para(doc,"[KONDISI KOSONG] Artikel belum tersedia — Admin dapat menambah artikel melalui Filament CMS.")
h2(doc,"5.2 Pola Detail Artikel")
for x in ["Breadcrumb: Beranda / Artikel / {judul}","Kategori","Judul","Ringkasan","Tanggal terbit dan penulis","Isi artikel","Terakhir diperbarui dan pemilik konten"]: bullet(doc,x)
articles=[("Panduan Cepat Menemukan Layanan Kampus","Panduan","Mulai dari kebutuhanmu, lalu ikuti satu tindakan utama pada setiap layanan.","Gunakan pencarian, pilih layanan, baca persyaratan, lalu ikuti tombol tindakan menuju kanal resmi."),("Kampus Inklusif: Mengenal Fasilitas Aksesibilitas","Aksesibilitas","Kenali dukungan akses dan cara meminta pendampingan sebelum berkunjung.","Hubungi kanal resmi sebelum kunjungan agar kebutuhan akses dapat disiapkan."),("Persiapan Administrasi Mahasiswa Baru","Mahasiswa Baru","Daftar ringkas untuk menyiapkan proses registrasi.","Pastikan data identitas konsisten dan proses dilakukan melalui portal resmi SMUP.")]
for i,a in enumerate(articles,1): h2(doc,f"5.{i+2} {a[0]}"); field(doc,"Kategori",a[1]); field(doc,"Ringkasan",a[2]); field(doc,"Isi",a[3]); field(doc,"Penulis","Tim ULT Unpad"); field(doc,"Pemilik konten","ULT Unpad")

page_break(doc); h1(doc,"6. FAQ")
label(doc,"Pengantar"); para(doc,"Pusat bantuan mandiri"); para(doc,"Pertanyaan yang sering diajukan"); field(doc,"Placeholder","Cari pertanyaan..."); field(doc,"Tombol","Cari FAQ")
faqs=[("Umum","Bagaimana cara menemukan layanan yang tepat?","Gunakan pencarian global atau pilih kategori berdasarkan kebutuhan. Setiap layanan memiliki langkah berikutnya yang jelas."),("Umum","Apakah website ULT membuat tiket pengaduan?","Tidak. Website ULT adalah portal informasi yang mengarahkan Anda ke sistem atau kanal resmi."),("Teknis","Apa yang dilakukan jika tautan eksternal bermasalah?","Buka halaman Kontak untuk kanal alternatif resmi, lalu sampaikan nama layanan dan kendalanya."),("Aksesibilitas","Bagaimana meminta pendampingan aksesibilitas?","Buka Layanan Disabilitas lalu hubungi ULT sebelum kunjungan.")]
for c,q,a in faqs: h2(doc,q); field(doc,"Kategori",c); para(doc,a)
para(doc,"[KONDISI KOSONG] Jawaban belum ditemukan — Coba kata kunci lain atau hubungi ULT.")

page_break(doc); h1(doc,"7. Kontak")
label(doc,"Pengantar"); para(doc,"Hubungi kami"); para(doc,"Kami bantu arahkan ke kanal yang tepat")
para(doc,"Sebelum menghubungi, coba cari layanan atau FAQ. Untuk kebutuhan lanjutan gunakan kontak resmi berikut.")
contacts=[("WhatsApp Admin ULT","+62 800-0000-0000","https://wa.me/6280000000000","[CONTOH/VERIFIKASI] Nomor contoh—ganti dengan nomor resmi melalui Filament."),("Helpdesk Unpad","helpdesk.unpad.ac.id","https://helpdesk.unpad.ac.id/","Kanal bantuan resmi Universitas Padjadjaran."),("Instagram ULT Unpad","@ult_unpad","https://www.instagram.com/ult_unpad?igsh=aThjeHo1YmlmcTl6","Ikuti informasi dan aktivitas terbaru ULT Unpad."),("TikTok ULT Unpad","@ult_unpad","https://www.tiktok.com/@ult_unpad","Temukan konten layanan dan informasi singkat ULT Unpad."),("Email ULT Unpad","ult@unpad.ac.id","mailto:ult@unpad.ac.id","[CONTOH/VERIFIKASI] Alamat contoh—verifikasi email resmi sebelum publikasi.")]
for n,v,u,d in contacts: h2(doc,n); field(doc,"Nilai",v); field(doc,"URL",u); para(doc,d)
h2(doc,"Kunjungan Langsung"); para(doc,"Pastikan membawa persyaratan yang tercantum pada halaman layanan dan periksa jam operasional terbaru.")
para(doc,"[KONDISI KOSONG] Kontak resmi belum tersedia — Admin dapat menambahkan WhatsApp, Instagram, TikTok, email, telepon, dan alamat melalui CMS.")

page_break(doc); h1(doc,"8. Pencarian")
label(doc,"Pengantar"); para(doc,"Pencarian global")
field(doc,"Judul tanpa kata kunci","Cari layanan dan informasi"); field(doc,"Judul dengan kata kunci","Hasil untuk “{kata kunci}”"); field(doc,"Placeholder","Masukkan minimal 2 karakter"); field(doc,"Tombol","Cari")
para(doc,"Hasil dikelompokkan menjadi Layanan, Artikel, dan FAQ. Setiap kelompok menampilkan jumlah hasil.")
para(doc,"[KONDISI KOSONG] Belum ada hasil yang cocok — Coba istilah lain, buka FAQ, atau hubungi ULT.")

page_break(doc); h1(doc,"9. Versi Bahasa Inggris")
para(doc,"Bagian ini mencatat semua salinan EN yang tersedia pada antarmuka. Konten CMS awal—layanan, artikel, FAQ, portal, dan kontak—masih menggunakan Bahasa Indonesia yang sama pada mode EN.")
h2(doc,"9.1 Global Interface")
for x in ["Skip to main content","Home","ULT Profile","Services","Articles","FAQ","Contact","Open search","Accessibility","Open menu","Language","What are you looking for?","Example: student card, tuition, diploma...","Search"]: bullet(doc,x)
h2(doc,"9.2 Home")
english_home=["One gateway to Unpad services","Find campus services without the confusion.","From academic needs and new student support to IT and inclusive services—we guide you to the right next step.","Search services, information, or questions...","Search now","Popular:","Start with your needs","I need a service for...","View all services","services available","Featured services","What students need most often","Explore all","Inclusive campus","Services designed to be accessible to everyone","ULT Unpad helps people with disabilities find facility information, online alternatives, and appropriate assistance channels.","Accessible locations and facilities","Online service alternatives","Screen-reader and keyboard friendly content","View disability services","Latest articles","Important information for campus life","View all articles","Quick answers","Frequently asked questions","Still need help? Search again or contact the appropriate official channel.","Read all FAQs","Continue to official systems","Unpad service portals & channels"]
for x in english_home: bullet(doc,x)
h2(doc,"9.3 Profile")
english_profile=["About ULT Unpad","A trusted gateway to Unpad services.","Since 2015, ULT Unpad has continued to grow as an integrated, accessible, and user-focused public service gateway.","Established; One Stop Service; Service values","History; Vision & Mission; PASTI; Legal Basis; Our People","Our journey","Growing from one roof into one integrated gateway","ULT Unpad was established in 2015 using a “one roof” service model in a modest room, supported by representatives assigned from different work units.","In line with Unpad’s institutional quality strategy and Continuing Quality Improvement principles, ULT continues to remove service barriers across work units.","One Roof Service — Service representatives from multiple work units served users together in one location.","Transformation into One Stop Service — ULT transformed into an integrated “one stop service”, managing requests from submission through completion in one place.","Hybrid & Inclusive Service — ULT combines in-person and online service while continuing to strengthen disability-friendly access.","Our direction; Vision and mission","Vision — To realize professional, accountable, simple, transparent, and informative public information services at Universitas Padjadjaran, fulfilling the public’s right to information in accordance with applicable laws.","Mission — Provide fast, timely, simple, and high-quality public services.","Mission — Develop effective and efficient technology-based information services.","Mission — Improve governance quality through optimized one-stop public service.","Our motto","Competent, responsive service; Responsible and measurable; Easy-to-understand processes; Clear requirements and procedures; Useful, relevant information.","Accountable service; Legal basis","Unpad established a legal foundation for the formation and management of ULT as part of its commitment to service-user satisfaction.","Service with empathy; Service officers; Prepared to provide excellent service","ULT officers pass a rigorous selection process, including psychological assessment and internships across Universitas Padjadjaran directorates.","They receive service-excellence competencies through regular training, enabling ULT to remain a reliable front line for lecturers, staff, students, and the public.","Contact ULT","Comfortable service area; An inclusive public space; A clear service gateway"]
for x in english_profile: bullet(doc,x)
h2(doc,"9.4 Contact")
for x in ["Contact us","We will guide you to the right channel","Search our services or FAQ first, then use one of the official contact channels below.","In-person visit","Bring all required documents and check the latest service hours before visiting."]: bullet(doc,x)
h2(doc,"9.5 Accessibility")
for x in ["Accessibility Menu (CTRL+U)","Adjust the website to your needs. Settings are saved automatically.","Increase Text","Decrease Text","High Contrast","Dark Mode","Highlight Links","Text Spacing","Pause Animation","Hide Images","Dyslexia Friendly","Large Cursor","Line Height","Text Alignment","Saturation","Reset All Accessibility Settings","Close accessibility menu"]: bullet(doc,x)
h2(doc,"9.6 Footer")
for x in ["An information and service-routing portal for the academic community and the public.","Explore","ULT Profile","Service Directory","Articles & Information","Frequently Asked Questions","Main Links","Contact ULT","Location","Contact information must be verified by an administrator before production.","Admin Login"]: bullet(doc,x)

page_break(doc); h1(doc,"10. Checklist Pemeriksaan Copywriting")
for x in ["Konsistensi istilah: ULT, Unpad, sivitas akademika, pengguna layanan, admin.","Konsistensi kapitalisasi judul, label tombol, dan nama sistem resmi.","Kejelasan sasaran pengguna dan langkah selanjutnya pada setiap layanan.","Verifikasi semua klaim kelembagaan, tanggal, dasar hukum, dan istilah PASTI.","Verifikasi nomor WhatsApp, alamat email, alamat lokasi, jam operasional, dan unit penanggung jawab.","Periksa penggunaan bahasa campuran: hybrid, routing, exchange, study permit, assessment/asesmen.","Pastikan salinan EN bukan hanya antarmuka, tetapi juga konten CMS bila website akan benar-benar bilingual.","Periksa nada bahasa: formal institusional, tetap mudah dipahami mahasiswa dan masyarakat umum.","Periksa apakah teks kondisi kosong dan pesan bantuan sudah mengarahkan pengguna dengan jelas.","Tandai konten yang perlu persetujuan legal, PPID, atau unit pemilik layanan."]: bullet(doc,x)

doc.core_properties.title="Inventaris Konten Website ULT Unpad"
doc.core_properties.subject="Dokumen pemeriksaan copywriting manual"
doc.core_properties.author="Unit Layanan Terpadu Universitas Padjadjaran"
doc.core_properties.keywords="ULT Unpad, copywriting, website, konten"
doc.save(OUT)
print(OUT)
