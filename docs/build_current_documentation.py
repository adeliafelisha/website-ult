from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DB = ROOT / "database" / "database.sqlite"
PURPLE = "6F1D69"
DARK = "292129"
PALE = "F7F0F6"
GRAY = "D9D9D9"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRAY)
        tbl_borders.append(node)


def configure(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for name, size in (("Title", 25), ("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)


def intro(doc, text, updated="4 September 2026"):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Subtitle"]
    p.runs[0].font.color.rgb = RGBColor(70, 60, 70)
    meta = doc.add_paragraph(f"Versi dokumentasi: {updated}")
    meta.runs[0].bold = True
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(111, 29, 105)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    borders(tbl)
    for index, header in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        cell.text = str(header)
        shade(cell, DARK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
    for row_index, values in enumerate(rows):
        cells = tbl.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = "" if value is None else str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                shade(cells[index], PALE)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for row in tbl.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return tbl


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_csv_guide(doc, name, path, required, important, rules):
    doc.add_heading(name, level=2)
    doc.add_paragraph(f"Template: docs/{path}")
    doc.add_paragraph("Kolom wajib: " + ", ".join(required) + ".")
    table(doc, ["Kolom", "Kegunaan"], important, [2.2, 4.4])
    for rule in rules:
        bullet(doc, rule)


def build_admin_guide():
    doc = Document()
    configure(doc, "Panduan Admin Website ULT Unpad")
    intro(doc, "Panduan operasional untuk mengelola konten, publikasi, data survei, tombol kontak, dan impor CSV melalui Filament CMS.")

    doc.add_heading("Akses dan kewenangan admin", level=1)
    doc.add_paragraph("Buka /admin pada domain website. Akun hanya dapat masuk jika berstatus admin, email sudah terverifikasi, dan domain email termasuk daftar yang diizinkan. Admin website mengelola konten publik; pengaturan server, database, backup, dan akun tetap menjadi tanggung jawab administrator teknis.")
    table(doc, ["Area", "Yang dapat dilakukan admin", "Dampak ke website"], [
        ["Dashboard", "Melihat ringkasan konten dan trafik bulanan", "Tidak mengubah konten"],
        ["Kategori Layanan", "Tambah, edit, urutkan, tandai unggulan", "Mengubah kelompok layanan"],
        ["Layanan", "CRUD, tombol kontak, publikasi, impor CSV", "Mengubah direktori dan detail layanan"],
        ["Artikel", "CRUD, gambar, link opsional, publikasi, impor CSV", "Mengubah artikel dan carousel beranda"],
        ["FAQ", "CRUD, kategori, link opsional, unggulan, impor CSV", "Mengubah FAQ dan FAQ beranda"],
        ["Survei Kepuasan", "Empat skor per tahun dan dua tautan", "Mengubah bagian SKM pada Profil"],
        ["Kontak", "WhatsApp, helpdesk, sosial, email, alamat", "Mengubah halaman kontak dan footer"],
        ["Tautan Cepat", "Nama, deskripsi, URL, urutan, status", "Mengubah tautan publik"],
    ], [1.25, 3.15, 2.25])

    doc.add_heading("Tema dashboard", level=1)
    doc.add_paragraph("Menu pengguna di kanan atas menyediakan Light, Dark, atau System. Pilihan ini hanya memengaruhi dashboard admin dan disimpan pada browser yang digunakan.")

    doc.add_heading("Alur kerja perubahan konten", level=1)
    for step in [
        "Buka resource yang akan dikelola dan cari record yang tepat.",
        "Buat atau edit record. Lengkapi Bahasa Indonesia dan English pada record yang sama.",
        "Periksa slug, URL resmi, pemilik konten, tanggal, dan status publikasi.",
        "Simpan sebagai draft dengan mematikan Terbit jika konten belum disetujui.",
        "Aktifkan Terbit atau Aktif setelah pemeriksaan selesai.",
        "Buka halaman publik dalam ID dan EN, lalu periksa tampilan desktop dan ponsel.",
    ]:
        numbered(doc, step)
    doc.add_paragraph("Frontend dan admin menggunakan database yang sama. Perubahan yang disimpan dan dipublikasikan tampil pada request berikutnya tanpa deploy ulang.")

    doc.add_heading("Operasi data umum", level=1)
    table(doc, ["Operasi", "Cara", "Catatan"], [
        ["Tambah", "Klik New atau Create, isi form, lalu Save", "Gunakan slug unik"],
        ["Edit", "Klik Edit pada baris yang dipilih", "Periksa kedua bahasa"],
        ["Hapus", "Klik Delete lalu konfirmasi", "Penghapusan tidak mudah dibatalkan"],
        ["Terbitkan massal", "Centang baris lalu pilih Terbitkan", "Tersedia pada layanan, artikel, FAQ, dan survei"],
        ["Batalkan terbit massal", "Centang baris lalu pilih Batalkan terbit", "Konten hilang dari halaman publik"],
        ["Cari dan filter", "Gunakan pencarian atau filter tabel", "Tidak mengubah data"],
    ], [1.35, 2.7, 2.6])

    doc.add_heading("Mengelola kategori dan layanan", level=1)
    doc.add_heading("Kategori layanan", level=2)
    doc.add_paragraph("Buat kategori lebih dahulu karena setiap layanan wajib terkait ke satu category_slug. Kategori yang masih mempunyai layanan tidak dapat dihapus. Kolom utama adalah nama ID/EN, slug, deskripsi ID/EN, urutan, dan status unggulan.")
    doc.add_heading("Layanan", level=2)
    doc.add_paragraph("Form layanan berisi identitas, konten ID/EN, persyaratan, dokumen, prosedur, informasi operasional, SEO, gambar, publikasi, dan tombol kontak.")
    bullet(doc, "delivery_type harus online, offline, atau hybrid.")
    bullet(doc, "Tombol kontak minimal satu. Default-nya Hubungi Admin ULT ke WhatsApp aktif.")
    bullet(doc, "Admin dapat mengubah label ID/EN, kanal, URL, urutan, serta menambahkan tombol Helpdesk, email, telepon, website, atau kanal lain.")
    bullet(doc, "Jika URL kontak global berubah, perbarui menu Kontak dan tombol layanan yang menyimpan URL lama.")

    doc.add_heading("Mengelola artikel dan FAQ", level=1)
    doc.add_heading("Artikel", level=2)
    doc.add_paragraph("Isi judul, kategori, ringkasan, konten, penulis, pemilik konten, gambar, deskripsi SEO, dan bahasa Inggris. external_url bersifat opsional; jika diisi, external_label menjadi tombol pada akhir artikel. Artikel unggulan dapat muncul pada carousel beranda.")
    doc.add_heading("FAQ", level=2)
    doc.add_paragraph("Isi pertanyaan, jawaban, kategori, sasaran, urutan, dan versi Inggris. external_url bersifat opsional; jika diisi, external_label menjadi tombol di bawah jawaban. is_featured menampilkan FAQ pada beranda.")

    doc.add_heading("Mengelola survei kepuasan", level=1)
    doc.add_paragraph("Satu record mewakili satu tahun dan tahun harus unik. Isi skor Triwulan 1 sampai 4 dalam rentang 0-100. source_url menghasilkan tombol Lihat data asli, sedangkan questionnaire_url menghasilkan tombol Kuesioner survei kepuasan masyarakat. Skor kosong ditampilkan sebagai data belum tersedia.")

    doc.add_heading("Mengelola kontak dan tautan cepat", level=1)
    doc.add_paragraph("Kontak mendukung email, phone, whatsapp, helpdesk, instagram, tiktok, address, dan other. Gunakan URL lengkap dengan https://. Tautan sosial aktif tampil di footer; semua kontak aktif tampil di halaman Kontak. Tautan Cepat berisi nama ID/EN, deskripsi ID/EN, URL, urutan, dan status aktif.")

    doc.add_heading("Impor dataset CSV", level=1)
    doc.add_paragraph("Impor tersedia untuk Layanan, Artikel, dan FAQ. Gunakan UTF-8, pemisah koma, satu baris header, dan satu record per baris. Jangan mengubah nama header. Simpan teks yang mengandung koma atau HTML di dalam tanda kutip ganda. Nilai boolean menerima 1, true, ya, yes, publish, atau published sebagai benar; gunakan 0 untuk draft.")
    add_csv_guide(doc, "CSV layanan", "service-import-template.csv", ["category_slug", "title", "summary"], [
        ["category_slug", "Slug kategori yang sudah tersedia di admin"],
        ["slug", "Identitas URL unik; jika kosong dibuat dari title"],
        ["delivery_type", "online, offline, atau hybrid"],
        ["contact_1_label sampai contact_3_label", "Tulisan tombol kontak Indonesia"],
        ["contact_1_label_en sampai contact_3_label_en", "Tulisan tombol kontak English"],
        ["contact_1_channel sampai contact_3_channel", "whatsapp, helpdesk, email, phone, website, atau other"],
        ["contact_1_url sampai contact_3_url", "URL lengkap tujuan tombol"],
        ["is_featured dan is_published", "1 untuk aktif, 0 untuk tidak aktif"],
    ], [
        "Buat kategori sebelum impor dan salin slug persis dari menu Kategori Layanan.",
        "Jika semua contact_N_url kosong, sistem otomatis membuat tombol WhatsApp Admin ULT.",
        "Impor dengan slug yang sudah ada akan memperbarui layanan tersebut.",
    ])
    add_csv_guide(doc, "CSV artikel", "article-import-template.csv", ["title", "category", "excerpt", "content"], [
        ["slug", "URL unik; jika kosong dibuat dari title"],
        ["content dan content_en", "HTML aman seperti p, ol, ul, li, strong, dan link"],
        ["external_url", "Opsional; menampilkan tombol eksternal"],
        ["external_label dan external_label_en", "Tulisan tombol eksternal"],
        ["seo_description", "Ringkasan untuk mesin pencari"],
        ["is_featured dan is_published", "1 untuk aktif, 0 untuk draft"],
    ], ["Impor dengan slug yang sama memperbarui artikel.", "Gambar utama tetap diunggah melalui form admin setelah impor."])
    add_csv_guide(doc, "CSV FAQ", "faq-import-template.csv", ["question", "answer", "category"], [
        ["question", "Kunci pembaruan record; harus konsisten"],
        ["answer", "Jawaban Indonesia; HTML sederhana diperbolehkan"],
        ["audience", "Sasaran seperti Mahasiswa atau Umum"],
        ["external_url", "Opsional; menampilkan tombol link"],
        ["sort_order", "Angka urutan; angka kecil tampil lebih dahulu"],
        ["is_featured dan is_published", "1 untuk aktif, 0 untuk tidak aktif"],
    ], ["Impor dengan question yang sama memperbarui FAQ.", "Periksa FAQ unggulan di beranda setelah impor."])

    doc.add_heading("Validasi sebelum impor", level=2)
    for item in [
        "Buka CSV di editor yang tidak mengubah format URL atau karakter UTF-8.",
        "Pastikan jumlah kolom setiap baris sama dengan header.",
        "Pastikan slug unik, category_slug valid, URL memakai https://, dan HTML memiliki tag penutup.",
        "Impor sebagai draft terlebih dahulu, periksa hasil, lalu terbitkan secara massal.",
        "Simpan salinan CSV sumber sebagai arsip perubahan.",
    ]:
        bullet(doc, item)

    doc.add_heading("Gambar dan media", level=1)
    doc.add_paragraph("Format gambar yang diterima adalah JPG, PNG, dan WebP dengan ukuran maksimal 5 MB. Gunakan nama deskriptif, kompres gambar, dan pastikan hak penggunaan. Pada server, php artisan storage:link harus dijalankan satu kali agar upload dapat dibuka publik.")

    doc.add_heading("Pemeriksaan sebelum publikasi", level=1)
    for item in ["Bahasa Indonesia dan English lengkap", "Judul dan ringkasan mudah dipahami", "Persyaratan dan prosedur telah diverifikasi", "URL resmi dan memakai HTTPS", "Tombol kontak dapat dibuka", "Hak penggunaan gambar jelas", "Pemilik konten terisi", "Status dan waktu terbit benar", "Tampilan ponsel dan desktop diperiksa"]:
        bullet(doc, item)

    doc.add_heading("Pemecahan masalah", level=1)
    table(doc, ["Masalah", "Pemeriksaan"], [
        ["Tidak dapat login", "Periksa email, password, is_admin, email_verified_at, dan ADMIN_ALLOWED_DOMAINS"],
        ["Perubahan belum terlihat", "Periksa status Terbit, waktu terbit, bahasa, lalu muat ulang halaman"],
        ["Gambar tidak tampil", "Periksa storage link, izin storage/app/public, dan file upload"],
        ["Impor CSV gagal", "Periksa header, encoding UTF-8, category_slug, kolom wajib, tanda kutip, dan jumlah kolom"],
        ["Admin lambat", "Gunakan APP_DEBUG=false, OPcache, php artisan optimize, dan resource database yang cukup"],
        ["Teks mode gelap tidak jelas", "Muat ulang aset terbaru dan hapus cache browser/CDN setelah deployment"],
    ], [2.0, 4.65])

    doc.add_heading("Keamanan operator", level=1)
    for item in ["Gunakan akun individual, password unik, dan password manager.", "Jangan membagikan kredensial atau menyimpan password pada komputer bersama.", "Jangan mengunggah data pribadi, rahasia, atau dokumen internal ke konten publik.", "Verifikasi domain tujuan sebelum menyimpan tautan.", "Lakukan backup database dan storage secara rutin serta uji proses pemulihan.", "Logout setelah menggunakan perangkat bersama."]:
        bullet(doc, item)

    doc.add_heading("Batas kewenangan dan eskalasi", level=1)
    doc.add_paragraph("Admin konten tidak perlu mengubah file aplikasi, konfigurasi .env, struktur database, DNS, SSL, antrean, atau konfigurasi server. Laporkan pekerjaan berikut kepada administrator teknis dengan menyertakan waktu kejadian, URL halaman, langkah yang dilakukan, dan tangkapan layar error.")
    table(doc, ["Kondisi", "Tindakan admin", "Tujuan eskalasi"], [
        ["Error 500 atau halaman tidak dapat dibuka", "Hentikan percobaan berulang dan catat URL serta waktu", "Administrator aplikasi"],
        ["Database atau upload tidak tersimpan", "Simpan salinan konten secara lokal dan laporkan", "Administrator database/server"],
        ["Domain, HTTPS, atau DNS bermasalah", "Jangan mengubah link menjadi alamat sementara", "Tim infrastruktur"],
        ["Diduga akun disalahgunakan", "Logout, ganti password melalui prosedur resmi, dan laporkan", "Administrator keamanan"],
        ["Konten memuat data sensitif", "Batalkan terbit dan jangan menyebarkan salinan", "Pemilik konten dan keamanan"],
    ], [2.15, 2.8, 1.75])

    doc.save(DOCS / "admin-guide.docx")


def fetch_inventory():
    if not DB.exists():
        return {}
    connection = sqlite3.connect(DB)
    connection.row_factory = sqlite3.Row
    queries = {
        "categories": "select name, slug, is_featured from service_categories order by sort_order, name",
        "services": "select title, slug, delivery_type, is_published from services order by title",
        "articles": "select title, slug, category, is_published from articles order by title",
        "faqs": "select question, category, is_featured, is_published from faqs order by category, sort_order",
        "contacts": "select label, type, is_published from contacts order by sort_order, label",
        "links": "select name, url, is_published from quick_links order by sort_order, name",
        "surveys": "select year, is_published from satisfaction_surveys order by year desc",
    }
    result = {}
    for key, query in queries.items():
        try:
            result[key] = [dict(row) for row in connection.execute(query).fetchall()]
        except sqlite3.OperationalError:
            result[key] = []
    connection.close()
    return result


def status(value):
    return "Terbit" if value else "Draft"


def build_inventory():
    data = fetch_inventory()
    doc = Document()
    configure(doc, "Inventori Konten Website ULT Unpad")
    intro(doc, "Daftar struktur halaman, sumber data, status konten, komponen interaktif, dan penanggung jawab pembaruan untuk audit konten website.")

    doc.add_heading("Ringkasan inventori", level=1)
    rows = [
        ["Kategori layanan", len(data.get("categories", [])), "Admin Kategori Layanan"],
        ["Layanan", len(data.get("services", [])), "Admin Layanan"],
        ["Artikel", len(data.get("articles", [])), "Admin Artikel"],
        ["FAQ", len(data.get("faqs", [])), "Admin FAQ"],
        ["Kontak", len(data.get("contacts", [])), "Admin Kontak"],
        ["Tautan cepat", len(data.get("links", [])), "Admin Tautan Cepat"],
        ["Tahun survei", len(data.get("surveys", [])), "Admin Survei Kepuasan"],
    ]
    table(doc, ["Jenis konten", "Jumlah saat dokumentasi dibuat", "Sumber pembaruan"], rows, [2.2, 2.0, 2.5])

    doc.add_heading("Halaman publik dan komponen", level=1)
    table(doc, ["Halaman", "URL", "Konten utama", "Sumber"], [
        ["Beranda", "/", "Hero, pencarian, kelompok layanan, profil singkat, carousel artikel, FAQ, tautan", "Kategori, Artikel, FAQ, Tautan Cepat"],
        ["Profil ULT", "/profil", "Profil, sejarah, visi, misi, PASTI, SKM, dasar hukum, petugas, galeri", "Template halaman dan Survei Kepuasan"],
        ["Direktori layanan", "/layanan", "Pencarian kata kunci dan kelompok layanan", "Kategori dan Layanan"],
        ["Kategori layanan", "/layanan/kategori/{slug}", "Submenu kategori dan daftar layanan", "Kategori dan Layanan"],
        ["Detail layanan", "/layanan/{slug}", "Syarat, dokumen, prosedur, informasi, tombol kontak", "Layanan dan Kontak"],
        ["Artikel", "/artikel", "Daftar artikel terbit", "Artikel"],
        ["Detail artikel", "/artikel/{slug}", "Isi, gambar, metadata, tombol eksternal opsional", "Artikel"],
        ["FAQ", "/faq", "Pencarian, topik, accordion, tombol eksternal opsional", "FAQ"],
        ["Kontak", "/kontak", "Kanal resmi, sosial, lokasi", "Kontak"],
        ["Pencarian", "/pencarian?q=", "Hasil layanan, artikel, dan FAQ", "Layanan, Artikel, FAQ"],
    ], [1.25, 1.45, 2.55, 1.65])

    doc.add_heading("Fitur lintas halaman", level=1)
    table(doc, ["Fitur", "Cakupan", "Catatan inventori"], [
        ["Bahasa ID dan EN", "Semua halaman publik", "Field English dengan fallback Indonesia"],
        ["Mode siang dan malam", "Semua halaman publik", "Tersimpan di browser; warna menggunakan token tema"],
        ["Menu aksesibilitas", "Semua halaman publik", "Ukuran teks, kontras, tema, link, spasi, gerak, gambar, font, kursor, tinggi baris, perataan, saturasi"],
        ["Scroll reveal", "Komponen utama", "Dinonaktifkan bila reduced motion aktif"],
        ["Scroll to top", "Semua halaman publik", "Muncul setelah halaman digulir"],
        ["Pelacakan trafik", "Halaman HTML publik", "Agregat bulanan tanpa menyimpan data pribadi mentah"],
        ["Pelacakan outbound", "Tautan eksternal", "Mencatat label, URL, dan halaman sumber"],
    ], [1.75, 1.65, 3.5])

    sections = [
        ("Kategori layanan", "categories", ["Nama", "Slug", "Unggulan"], lambda r: [r["name"], r["slug"], "Ya" if r["is_featured"] else "Tidak"]),
        ("Layanan", "services", ["Judul", "Slug", "Jenis", "Status"], lambda r: [r["title"], r["slug"], r["delivery_type"], status(r["is_published"])]),
        ("Artikel", "articles", ["Judul", "Slug", "Kategori", "Status"], lambda r: [r["title"], r["slug"], r["category"], status(r["is_published"])]),
        ("FAQ", "faqs", ["Pertanyaan", "Kategori", "Beranda", "Status"], lambda r: [r["question"], r["category"], "Ya" if r["is_featured"] else "Tidak", status(r["is_published"])]),
        ("Kontak", "contacts", ["Label", "Jenis", "Status"], lambda r: [r["label"], r["type"], status(r["is_published"])]),
        ("Tautan cepat", "links", ["Nama", "URL", "Status"], lambda r: [r["name"], r["url"], status(r["is_published"])]),
        ("Survei kepuasan", "surveys", ["Tahun", "Status"], lambda r: [r["year"], status(r["is_published"])]),
    ]
    for heading, key, headers, mapper in sections:
        doc.add_heading(heading, level=1)
        records = data.get(key, [])
        if records:
            table(doc, headers, [mapper(record) for record in records])
        else:
            doc.add_paragraph("Belum ada record pada database lokal saat dokumentasi dibuat.")

    doc.add_heading("Kepemilikan dan siklus tinjauan", level=1)
    table(doc, ["Konten", "Tinjauan minimum", "Hal yang diperiksa"], [
        ["Layanan", "Bulanan atau saat kebijakan berubah", "Syarat, prosedur, jam, biaya, unit, tombol kontak"],
        ["Artikel", "Per kuartal", "Akurasi, tanggal, link sumber, hak gambar, bilingual"],
        ["FAQ", "Bulanan", "Pertanyaan berulang, jawaban, kategori, link"],
        ["Kontak", "Bulanan", "Nomor, akun sosial, URL, jam layanan"],
        ["Survei", "Setiap hasil triwulan", "Skor, tahun, sumber asli, kuesioner"],
        ["Konten statis Profil", "Semesteran", "Sejarah, visi, misi, dasar hukum, foto"],
    ], [1.55, 2.05, 3.25])

    doc.add_heading("Celah pengelolaan yang perlu dipantau", level=1)
    for item in [
        "Konten statis Profil masih dikelola melalui kode dan belum menjadi resource CMS.",
        "Gambar artikel tidak dapat diimpor lewat CSV dan perlu diunggah dari form admin.",
        "Perubahan URL WhatsApp global tidak otomatis mengganti URL yang sudah disimpan pada tombol tiap layanan.",
        "Persetujuan editorial bertingkat dan riwayat revisi konten belum tersedia.",
        "Pengujian aksesibilitas otomatis lengkap seperti axe belum menjadi bagian pipeline deployment.",
    ]:
        bullet(doc, item)

    doc.save(DOCS / "content-inventory.docx")


if __name__ == "__main__":
    build_admin_guide()
    build_inventory()
