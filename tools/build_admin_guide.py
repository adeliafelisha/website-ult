from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
d=Document(); s=d.sections[0]; s.top_margin=Inches(.8); s.bottom_margin=Inches(.8); s.left_margin=Inches(1); s.right_margin=Inches(1)
for st in ['Normal','Heading 1','Heading 2','Heading 3']:
 x=d.styles[st]; x.font.name='Calibri'; x._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); x._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
d.styles['Normal'].font.size=Pt(11); d.styles['Normal'].paragraph_format.space_after=Pt(6); d.styles['Normal'].paragraph_format.line_spacing=1.2
for st,sz in [('Heading 1',17),('Heading 2',14),('Heading 3',12)]: d.styles[st].font.size=Pt(sz); d.styles[st].font.color.rgb=RGBColor(112,29,107); d.styles[st].font.bold=True
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90); r=p.add_run('PANDUAN ADMIN'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor(112,29,107)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Website Unit Layanan Terpadu Unpad').font.size=Pt(16)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30); p.add_run('Filament CMS • Tambah, lihat, edit, hapus, publikasi, dan konten bilingual')
d.add_page_break()
def h(t,l=1): d.add_heading(t,level=l)
def p(t): d.add_paragraph(t)
def b(t): d.add_paragraph(t,style='List Bullet')
def n(t): d.add_paragraph(t,style='List Number')
h('1. Mengakses Admin'); p('Buka /admin pada domain website, misalnya https://ult.unpad.ac.id/admin. Masuk memakai akun admin resmi yang sudah diaktifkan. Akun harus berstatus admin, email terverifikasi, dan menggunakan domain email yang diizinkan.')
p('Jika kredensial ditolak, jangan membuat akun baru sembarangan. Hubungi administrator teknis untuk memeriksa record users, is_admin, email_verified_at, dan domain email.')
h('2. Tampilan dan Tema'); p('Gunakan menu pengguna di kanan atas untuk memilih tema Light, Dark, atau mengikuti sistem. Pilihan disimpan di browser dan tidak mengubah tampilan website publik.')
h('3. Cara Kerja Perubahan Konten'); p('Frontend dan admin menggunakan database yang sama. Setelah record disimpan dan status Terbit/Aktif dinyalakan, perubahan langsung dibaca website pada request berikutnya. Tidak perlu deploy ulang. Jika cache produksi digunakan dan teks belum berubah, jalankan php artisan optimize:clear lalu muat ulang halaman.')
h('4. Operasi CRUD Umum');
for x in ['Tambah: buka menu resource, klik New/Create, lengkapi form, lalu simpan.','Lihat: gunakan tabel, pencarian, filter, dan pagination.','Edit: klik Edit pada baris, ubah field, lalu Save.','Hapus: klik Delete pada baris dan konfirmasi. Untuk banyak data, centang baris lalu gunakan bulk delete.','Publikasi: nyalakan Terbit/Aktif. Untuk artikel dan layanan, atur waktu terbit bila diperlukan.','Draft: matikan Terbit agar konten tidak tampil publik.']: b(x)
h('5. Mengelola Bahasa Indonesia dan Inggris'); p('Isi tab Indonesia dan English pada record yang sama. Website memilih field English ketika pengguna memilih EN. Jika field English kosong, sistem memakai teks Indonesia sebagai fallback agar halaman tidak rusak. Untuk website bilingual yang baik, jangan biarkan tab English kosong.')
h('6. Menu Admin');
for title,text in [('Kategori Layanan','Mengatur kelompok layanan. Kategori yang masih memiliki layanan tidak dapat dihapus untuk menjaga integritas data.'),('Layanan','Mengatur judul, ringkasan, sasaran, syarat, dokumen, prosedur, CTA, URL, lokasi, jam, estimasi, biaya, unit, gambar, SEO, dan publikasi.'),('Artikel','Mengatur artikel ID/EN, kategori, ringkasan, isi, penulis, gambar, SEO, unggulan, dan waktu terbit.'),('FAQ','Mengatur pertanyaan dan jawaban ID/EN, kategori, urutan, unggulan, dan status terbit.'),('Kontak','Mengatur WhatsApp, helpdesk, Instagram, TikTok, email, telepon, alamat, URL, deskripsi, dan urutan.'),('Tautan Cepat','Mengatur portal resmi beserta nama/deskripsi ID/EN dan URL.')]: h(title,2); p(text)
h('7. Alur Kerja yang Disarankan');
for x in ['Buat atau edit konten sebagai draft.','Lengkapi Bahasa Indonesia dan English.','Periksa URL, persyaratan, unit pemilik, tanggal, dan kontak.','Gunakan Preview publik bila diperlukan.','Aktifkan Terbit/Aktif.','Buka halaman publik ID dan EN untuk memastikan hasil.','Catat perubahan penting dan jadwalkan tinjauan berkala.']: n(x)
h('8. Upload Gambar'); p('Format yang diterima: JPG, PNG, dan WebP, maksimal 5 MB. Gunakan nama file deskriptif, kompres gambar, dan pastikan hak penggunaan/copyright. Jalankan php artisan storage:link satu kali pada server agar upload tampil publik.')
h('9. Troubleshooting');
for x in ['Admin lambat: pastikan APP_DEBUG=false, OPcache aktif, database memiliki resource cukup, lalu jalankan php artisan optimize dan php artisan icons:cache.','Perubahan belum tampil: pastikan Terbit/Aktif menyala, waktu terbit tidak berada di masa depan, lalu bersihkan cache.','Gambar tidak tampil: periksa storage link dan izin tulis storage/app/public.','Tidak bisa menghapus kategori: pindahkan atau hapus layanan yang masih terkait terlebih dahulu.','Halaman EN masih Indonesia: lengkapi tab English pada record terkait.','Error 403 admin: periksa is_admin, email_verified_at, dan ADMIN_ALLOWED_DOMAINS.']: b(x)
h('10. Keamanan Operator');
for x in ['Jangan membagikan akun admin.','Gunakan password unik dan panjang serta password manager.','Logout pada komputer bersama.','Jangan mengunggah data pribadi atau dokumen rahasia.','Verifikasi link sebelum dipublikasikan.','Lakukan backup database dan storage secara rutin.']: b(x)
h('11. Checklist Sebelum Publikasi');
for x in ['Bahasa ID dan EN lengkap','Judul dan ringkasan jelas','Syarat/prosedur telah diverifikasi','URL resmi dan HTTPS','Kontak resmi','Gambar legal dan terkompresi','Pemilik konten terisi','Status terbit benar','Tampilan mobile diperiksa']: b('☐ '+x)
out=Path(__file__).resolve().parents[1]/'docs'/'admin-guide.docx'; out.parent.mkdir(exist_ok=True); d.core_properties.title='Panduan Admin Website ULT Unpad'; d.core_properties.author='ULT Unpad'; d.save(out); print(out)
